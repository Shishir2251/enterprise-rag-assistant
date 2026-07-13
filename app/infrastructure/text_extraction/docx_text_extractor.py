from pathlib import Path

from docx import Document

from app.business.dtos.extracted_text_dto import (
    ExtractedDocument,
    ExtractedPage,
)
from app.business.interfaces.text_extractor_interface import ITextExtractor


class DocxTextExtractor(ITextExtractor):

    def extract(self, file_path: Path) -> ExtractedDocument:
        document = Document(str(file_path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n\n".join(paragraphs)

        pages = []

        if text:
            pages.append(
                ExtractedPage(
                    page_number=None,
                    content=text,
                )
            )

        return ExtractedDocument(pages=pages)